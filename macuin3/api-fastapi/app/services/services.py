from sqlalchemy.orm import Session
from sqlalchemy import func, desc
from fastapi import HTTPException
from decimal import Decimal
from datetime import datetime
import io

from app.models.models import (
    Usuario, Rol, Cliente, Empleado, Categoria, Autoparte,
    Inventario, MovimientoInventario, Pedido, DetallePedido, HistorialEstatusPedido
)
from app.core.security import hash_password, verify_password, create_access_token


def authenticate_user(db: Session, email: str, password: str):
    user = db.query(Usuario).filter(Usuario.email == email).first()
    if not user or not verify_password(password, user.password_hash):
        raise HTTPException(status_code=401, detail="Credenciales inválidas")
    if not user.activo:
        raise HTTPException(status_code=403, detail="Cuenta desactivada")
    token = create_access_token({"sub": str(user.id), "rol": user.rol.nombre})
    return {"access_token": token, "token_type": "bearer", "user": _user_dict(user)}


def register_cliente(db: Session, data):
    if db.query(Usuario).filter(Usuario.email == data.email).first():
        raise HTTPException(status_code=400, detail="El email ya está registrado")
    rol = db.query(Rol).filter(Rol.nombre == "cliente_externo").first()
    if not rol:
        raise HTTPException(status_code=500, detail="Rol cliente_externo no encontrado")
    usuario = Usuario(email=data.email, password_hash=hash_password(data.password),
                      nombre=data.nombre, apellido=data.apellido, telefono=data.telefono, rol_id=rol.id, activo=True)
    db.add(usuario)
    db.flush()
    cliente = Cliente(usuario_id=usuario.id, razon_social=data.razon_social, rfc=data.rfc,
                      direccion=data.direccion, ciudad=data.ciudad, estado=data.estado,
                      codigo_postal=data.codigo_postal, tipo_cliente=data.tipo_cliente or "general")
    db.add(cliente)
    db.commit()
    db.refresh(usuario)
    token = create_access_token({"sub": str(usuario.id), "rol": rol.nombre})
    return {"access_token": token, "token_type": "bearer", "user": _user_dict(usuario)}


def _user_dict(user):
    return {"id": user.id, "email": user.email, "nombre": user.nombre, "apellido": user.apellido,
            "telefono": user.telefono, "rol": {"id": user.rol.id, "nombre": user.rol.nombre}, "activo": user.activo}


# ===== CATEGORIAS =====
def get_categorias(db: Session):
    return [{"id": c.id, "nombre": c.nombre, "descripcion": c.descripcion, "activo": c.activo}
            for c in db.query(Categoria).filter(Categoria.activo == True).all()]

def create_categoria(db: Session, data):
    cat = Categoria(nombre=data.nombre, descripcion=data.descripcion)
    db.add(cat); db.commit(); db.refresh(cat)
    return {"id": cat.id, "nombre": cat.nombre, "descripcion": cat.descripcion, "activo": cat.activo}

def delete_categoria(db: Session, cid: int):
    cat = db.query(Categoria).filter(Categoria.id == cid).first()
    if not cat: raise HTTPException(404, "Categoría no encontrada")
    cat.activo = False; db.commit()
    return {"message": "Categoría desactivada"}


# ===== AUTOPARTES =====
def _autoparte_dict(db, a):
    inv = db.query(Inventario).filter(Inventario.autoparte_id == a.id).first()
    return {"id": a.id, "sku": a.sku, "nombre": a.nombre, "descripcion": a.descripcion,
            "marca": a.marca, "categoria_id": a.categoria_id,
            "categoria_nombre": a.categoria.nombre if a.categoria else None,
            "precio": float(a.precio), "compatibilidad_vehicular": a.compatibilidad_vehicular,
            "imagen_url": a.imagen_url, "activo": a.activo,
            "stock_actual": inv.stock_actual if inv else 0, "stock_minimo": inv.stock_minimo if inv else 0}

def get_autopartes(db: Session, search=None, categoria_id=None, marca=None, solo_activos=True):
    q = db.query(Autoparte)
    if solo_activos: q = q.filter(Autoparte.activo == True)
    if search:
        q = q.filter((Autoparte.nombre.ilike(f"%{search}%")) | (Autoparte.sku.ilike(f"%{search}%")) |
                      (Autoparte.descripcion.ilike(f"%{search}%")) | (Autoparte.compatibilidad_vehicular.ilike(f"%{search}%")))
    if categoria_id: q = q.filter(Autoparte.categoria_id == categoria_id)
    if marca: q = q.filter(Autoparte.marca.ilike(f"%{marca}%"))
    return [_autoparte_dict(db, a) for a in q.order_by(Autoparte.nombre).all()]

def get_autoparte(db: Session, aid: int):
    a = db.query(Autoparte).filter(Autoparte.id == aid).first()
    if not a: raise HTTPException(404, "Autoparte no encontrada")
    return _autoparte_dict(db, a)

def create_autoparte(db: Session, data):
    if db.query(Autoparte).filter(Autoparte.sku == data.sku).first():
        raise HTTPException(400, "SKU ya existe")
    a = Autoparte(sku=data.sku, nombre=data.nombre, descripcion=data.descripcion, marca=data.marca,
                  categoria_id=data.categoria_id, precio=data.precio, compatibilidad_vehicular=data.compatibilidad_vehicular,
                  imagen_url=getattr(data, "imagen_url", None), activo=data.activo)
    db.add(a); db.flush()
    inv = Inventario(autoparte_id=a.id, stock_actual=data.stock_inicial, stock_minimo=data.stock_minimo)
    db.add(inv); db.commit(); db.refresh(a)
    return get_autoparte(db, a.id)

def update_autoparte(db: Session, aid: int, data):
    a = db.query(Autoparte).filter(Autoparte.id == aid).first()
    if not a: raise HTTPException(404, "Autoparte no encontrada")
    for k, v in data.model_dump(exclude_unset=True).items():
        setattr(a, k, v)
    db.commit(); db.refresh(a)
    return get_autoparte(db, a.id)

def delete_autoparte(db: Session, aid: int):
    a = db.query(Autoparte).filter(Autoparte.id == aid).first()
    if not a: raise HTTPException(404, "Autoparte no encontrada")
    a.activo = False; db.commit()
    return {"message": "Autoparte desactivada"}


# ===== INVENTARIO =====
def get_inventario(db: Session):
    return [{"id": inv.id, "autoparte_id": inv.autoparte_id,
             "autoparte_nombre": inv.autoparte.nombre if inv.autoparte else None,
             "autoparte_sku": inv.autoparte.sku if inv.autoparte else None,
             "stock_actual": inv.stock_actual, "stock_minimo": inv.stock_minimo,
             "ubicacion_almacen": inv.ubicacion_almacen}
            for inv in db.query(Inventario).all()]

def update_inventario(db: Session, autoparte_id: int, data, usuario_id: int):
    inv = db.query(Inventario).filter(Inventario.autoparte_id == autoparte_id).first()
    if not inv: raise HTTPException(404, "Inventario no encontrado")
    old = inv.stock_actual
    if data.tipo == "entrada": inv.stock_actual += data.cantidad
    elif data.tipo == "salida":
        if inv.stock_actual < data.cantidad: raise HTTPException(400, "Stock insuficiente")
        inv.stock_actual -= data.cantidad
    else: inv.stock_actual = data.cantidad
    db.add(MovimientoInventario(autoparte_id=autoparte_id, tipo=data.tipo, cantidad=data.cantidad,
                                 stock_anterior=old, stock_nuevo=inv.stock_actual, motivo=data.motivo, usuario_id=usuario_id))
    db.commit()
    return {"stock_actual": inv.stock_actual, "message": "Inventario actualizado"}


# ===== EMPLEADOS =====
def get_empleados(db: Session):
    emps = db.query(Empleado).all()
    result = []
    for e in emps:
        u = e.usuario
        result.append({"id": e.id, "usuario_id": u.id, "email": u.email, "nombre": u.nombre, "apellido": u.apellido,
                        "telefono": u.telefono, "numero_empleado": e.numero_empleado, "departamento": e.departamento,
                        "puesto": e.puesto, "activo": u.activo})
    return result

def create_empleado(db: Session, data):
    if db.query(Usuario).filter(Usuario.email == data.email).first():
        raise HTTPException(400, "Email ya registrado")
    if db.query(Empleado).filter(Empleado.numero_empleado == data.numero_empleado).first():
        raise HTTPException(400, "Número de empleado ya existe")
    rol = db.query(Rol).filter(Rol.nombre == "personal_interno").first()
    u = Usuario(email=data.email, password_hash=hash_password(data.password), nombre=data.nombre,
                apellido=data.apellido, telefono=data.telefono, rol_id=rol.id, activo=True)
    db.add(u); db.flush()
    e = Empleado(usuario_id=u.id, numero_empleado=data.numero_empleado, departamento=data.departamento, puesto=data.puesto)
    db.add(e); db.commit(); db.refresh(e)
    return get_empleados(db)[-1]

def update_empleado(db: Session, eid: int, data):
    e = db.query(Empleado).filter(Empleado.id == eid).first()
    if not e: raise HTTPException(404, "Empleado no encontrado")
    u = e.usuario
    d = data.model_dump(exclude_unset=True)
    if "nombre" in d: u.nombre = d["nombre"]
    if "apellido" in d: u.apellido = d["apellido"]
    if "telefono" in d: u.telefono = d["telefono"]
    if "departamento" in d: e.departamento = d["departamento"]
    if "puesto" in d: e.puesto = d["puesto"]
    if "activo" in d: u.activo = d["activo"]
    db.commit()
    return {"message": "Empleado actualizado"}

def delete_empleado(db: Session, eid: int):
    e = db.query(Empleado).filter(Empleado.id == eid).first()
    if not e: raise HTTPException(404, "Empleado no encontrado")
    e.usuario.activo = False; db.commit()
    return {"message": "Empleado desactivado"}


# ===== PEDIDOS =====
def _gen_folio(db):
    y = datetime.now().year
    last = db.query(Pedido).filter(Pedido.folio.like(f"PED-{y}-%")).order_by(desc(Pedido.id)).first()
    n = int(last.folio.split("-")[-1]) + 1 if last else 1
    return f"PED-{y}-{n:04d}"

def create_pedido(db: Session, data, usuario_id: int):
    cliente = db.query(Cliente).filter(Cliente.usuario_id == usuario_id).first()
    if not cliente: raise HTTPException(400, "Perfil de cliente no encontrado")
    if not data.items: raise HTTPException(400, "Agrega al menos un producto")
    folio = _gen_folio(db)
    subtotal = Decimal("0"); detalles = []
    for item in data.items:
        ap = db.query(Autoparte).filter(Autoparte.id == item.autoparte_id, Autoparte.activo == True).first()
        if not ap: raise HTTPException(404, f"Autoparte {item.autoparte_id} no encontrada")
        inv = db.query(Inventario).filter(Inventario.autoparte_id == item.autoparte_id).first()
        if not inv or inv.stock_actual < item.cantidad: raise HTTPException(400, f"Stock insuficiente: {ap.nombre}")
        s = ap.precio * item.cantidad; subtotal += s
        detalles.append({"autoparte_id": item.autoparte_id, "cantidad": item.cantidad, "precio_unitario": ap.precio, "subtotal": s})
        inv.stock_actual -= item.cantidad
        db.add(MovimientoInventario(autoparte_id=item.autoparte_id, tipo="salida", cantidad=item.cantidad,
                                     stock_anterior=inv.stock_actual + item.cantidad, stock_nuevo=inv.stock_actual,
                                     motivo=f"Pedido {folio}", usuario_id=usuario_id))
    imp = subtotal * Decimal("0.16"); total = subtotal + imp
    pedido = Pedido(folio=folio, cliente_id=cliente.id, estatus="recibido", subtotal=subtotal, impuesto=imp, total=total, notas=data.notas)
    db.add(pedido); db.flush()
    for d in detalles: db.add(DetallePedido(pedido_id=pedido.id, **d))
    db.add(HistorialEstatusPedido(pedido_id=pedido.id, estatus_nuevo="recibido", comentario="Pedido creado", usuario_id=usuario_id))
    db.commit()
    return _pedido_dict(db, pedido)

def get_pedidos(db, cliente_id=None, estatus=None):
    q = db.query(Pedido)
    if cliente_id: q = q.filter(Pedido.cliente_id == cliente_id)
    if estatus: q = q.filter(Pedido.estatus == estatus)
    return [_pedido_dict(db, p, False) for p in q.order_by(desc(Pedido.created_at)).all()]

def get_pedido(db, pid):
    p = db.query(Pedido).filter(Pedido.id == pid).first()
    if not p: raise HTTPException(404, "Pedido no encontrado")
    return _pedido_dict(db, p)

def get_pedidos_cliente(db, usuario_id):
    c = db.query(Cliente).filter(Cliente.usuario_id == usuario_id).first()
    if not c: raise HTTPException(400, "Cliente no encontrado")
    return get_pedidos(db, cliente_id=c.id)

def cambiar_estatus(db, pid, nuevo, comentario, uid):
    p = db.query(Pedido).filter(Pedido.id == pid).first()
    if not p: raise HTTPException(404, "Pedido no encontrado")
    trans = {"recibido": ["en_proceso", "cancelado"], "en_proceso": ["enviado", "cancelado"],
             "enviado": ["entregado"], "entregado": [], "cancelado": []}
    if nuevo not in trans.get(p.estatus, []): raise HTTPException(400, f"Transición inválida: {p.estatus} → {nuevo}")
    if nuevo == "cancelado":
        for d in p.detalles:
            inv = db.query(Inventario).filter(Inventario.autoparte_id == d.autoparte_id).first()
            if inv:
                old = inv.stock_actual; inv.stock_actual += d.cantidad
                db.add(MovimientoInventario(autoparte_id=d.autoparte_id, tipo="entrada", cantidad=d.cantidad,
                                             stock_anterior=old, stock_nuevo=inv.stock_actual,
                                             motivo=f"Cancelación {p.folio}", usuario_id=uid))
    ant = p.estatus; p.estatus = nuevo
    db.add(HistorialEstatusPedido(pedido_id=p.id, estatus_anterior=ant, estatus_nuevo=nuevo, comentario=comentario, usuario_id=uid))
    db.commit()
    return _pedido_dict(db, p)

def cancelar_pedido_cliente(db, pid, uid):
    c = db.query(Cliente).filter(Cliente.usuario_id == uid).first()
    p = db.query(Pedido).filter(Pedido.id == pid).first()
    if not p or p.cliente_id != c.id: raise HTTPException(404, "Pedido no encontrado")
    if p.estatus != "recibido": raise HTTPException(400, "Solo se pueden cancelar pedidos 'recibido'")
    return cambiar_estatus(db, pid, "cancelado", "Cancelado por cliente", uid)

def _pedido_dict(db, p, details=True):
    c = p.cliente; u = c.usuario if c else None
    r = {"id": p.id, "folio": p.folio, "estatus": p.estatus, "subtotal": float(p.subtotal),
         "impuesto": float(p.impuesto), "total": float(p.total), "notas": p.notas,
         "created_at": p.created_at.isoformat() if p.created_at else None,
         "updated_at": p.updated_at.isoformat() if p.updated_at else None,
         "cliente": {"id": c.id, "razon_social": c.razon_social, "nombre": u.nombre if u else None,
                     "apellido": u.apellido if u else None, "email": u.email if u else None} if c else None}
    if details:
        r["detalles"] = [{"id": d.id, "autoparte_id": d.autoparte_id,
                          "autoparte_nombre": d.autoparte.nombre if d.autoparte else None,
                          "autoparte_sku": d.autoparte.sku if d.autoparte else None,
                          "cantidad": d.cantidad, "precio_unitario": float(d.precio_unitario), "subtotal": float(d.subtotal)} for d in p.detalles]
        r["historial"] = [{"id": h.id, "estatus_anterior": h.estatus_anterior, "estatus_nuevo": h.estatus_nuevo,
                           "comentario": h.comentario, "usuario_nombre": f"{h.usuario.nombre} {h.usuario.apellido}" if h.usuario else None,
                           "created_at": h.created_at.isoformat() if h.created_at else None} for h in p.historial]
    return r


# ===== DOCX PEDIDO INDIVIDUAL =====
def generar_docx_pedido(db, pid):
    from docx import Document
    from docx.shared import Pt, RGBColor
    p = db.query(Pedido).filter(Pedido.id == pid).first()
    if not p: raise HTTPException(404, "Pedido no encontrado")
    doc = Document()
    doc.add_heading("MACUIN - Autopartes Automotrices", 0)
    doc.add_heading(f"Pedido: {p.folio}", level=1)
    doc.add_paragraph(f"Estatus: {p.estatus.upper()}    Fecha: {p.created_at.strftime('%d/%m/%Y %H:%M') if p.created_at else ''}")
    if p.cliente and p.cliente.usuario:
        doc.add_paragraph(f"Cliente: {p.cliente.razon_social or ''} — {p.cliente.usuario.nombre} {p.cliente.usuario.apellido}")
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdrs = table.rows[0].cells
    for i, h in enumerate(["SKU", "Producto", "Cant.", "P.Unit.", "Subtotal"]):
        hdrs[i].text = h
        run = hdrs[i].paragraphs[0].runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for d in p.detalles:
        row = table.add_row().cells
        row[0].text = d.autoparte.sku if d.autoparte else ""
        row[1].text = d.autoparte.nombre if d.autoparte else ""
        row[2].text = str(d.cantidad)
        row[3].text = f"${d.precio_unitario:,.2f}"
        row[4].text = f"${d.subtotal:,.2f}"
    doc.add_paragraph("")
    doc.add_paragraph(f"Subtotal: ${p.subtotal:,.2f}")
    doc.add_paragraph(f"IVA 16%:  ${p.impuesto:,.2f}")
    t = doc.add_paragraph(f"TOTAL:    ${p.total:,.2f}")
    t.runs[0].bold = True
    if p.notas: doc.add_paragraph(f"Notas: {p.notas}")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ===== REPORTES =====
def get_reporte_ventas(db):
    tv = db.query(func.sum(Pedido.total)).filter(Pedido.estatus != "cancelado").scalar() or 0
    tp = db.query(func.count(Pedido.id)).scalar() or 0
    est = {e: c for e, c in db.query(Pedido.estatus, func.count(Pedido.id)).group_by(Pedido.estatus).all()}
    ct = [{"id": c[0], "razon_social": c[1], "total_pedidos": c[2], "monto_total": float(c[3] or 0)}
          for c in db.query(Cliente.id, Cliente.razon_social, func.count(Pedido.id).label("tp"),
                            func.sum(Pedido.total).label("mt")).join(Pedido).filter(Pedido.estatus != "cancelado")
          .group_by(Cliente.id, Cliente.razon_social).order_by(desc("tp")).limit(10).all()]
    pt = [{"id": p[0], "nombre": p[1], "sku": p[2], "total_vendido": int(p[3] or 0)}
          for p in db.query(Autoparte.id, Autoparte.nombre, Autoparte.sku,
                            func.sum(DetallePedido.cantidad).label("tv")).join(DetallePedido).join(Pedido)
          .filter(Pedido.estatus != "cancelado").group_by(Autoparte.id, Autoparte.nombre, Autoparte.sku)
          .order_by(desc("tv")).limit(10).all()]
    sb = [{"id": inv.autoparte.id, "nombre": inv.autoparte.nombre, "sku": inv.autoparte.sku,
           "stock_actual": inv.stock_actual, "stock_minimo": inv.stock_minimo}
          for inv in db.query(Inventario).filter(Inventario.stock_actual <= Inventario.stock_minimo).all() if inv.autoparte]
    return {"total_ventas": float(tv), "total_pedidos": tp, "pedidos_por_estatus": est,
            "clientes_top": ct, "productos_top": pt, "productos_stock_bajo": sb}


def get_reporte_inventario(db):
    rows = []
    for inv in db.query(Inventario).all():
        if inv.autoparte:
            rows.append({"sku": inv.autoparte.sku, "nombre": inv.autoparte.nombre,
                         "marca": inv.autoparte.marca, "categoria": inv.autoparte.categoria.nombre if inv.autoparte.categoria else "",
                         "stock_actual": inv.stock_actual, "stock_minimo": inv.stock_minimo,
                         "ubicacion": inv.ubicacion_almacen or "",
                         "estatus": "BAJO" if inv.stock_actual <= inv.stock_minimo else "OK"})
    return {"total_productos": len(rows), "productos_stock_bajo": sum(1 for r in rows if r["estatus"] == "BAJO"), "detalle": rows}


def get_reporte_pedidos(db):
    pedidos = db.query(Pedido).order_by(desc(Pedido.created_at)).all()
    rows = []
    for p in pedidos:
        c = p.cliente; u = c.usuario if c else None
        rows.append({"folio": p.folio, "cliente": c.razon_social or (f"{u.nombre} {u.apellido}" if u else ""),
                     "estatus": p.estatus, "subtotal": float(p.subtotal), "impuesto": float(p.impuesto),
                     "total": float(p.total), "fecha": p.created_at.strftime("%d/%m/%Y") if p.created_at else ""})
    return {"total_pedidos": len(rows), "monto_total": sum(r["total"] for r in rows), "detalle": rows}


def get_reporte_clientes(db):
    rows = []
    for c in db.query(Cliente).all():
        u = c.usuario
        pedidos = db.query(Pedido).filter(Pedido.cliente_id == c.id, Pedido.estatus != "cancelado").all()
        rows.append({"nombre": f"{u.nombre} {u.apellido}" if u else "", "email": u.email if u else "",
                     "razon_social": c.razon_social or "", "rfc": c.rfc or "", "ciudad": c.ciudad or "",
                     "estado": c.estado or "", "total_pedidos": len(pedidos),
                     "monto_total": float(sum(p.total for p in pedidos))})
    return {"total_clientes": len(rows), "detalle": rows}


# ===== GENERADORES PDF REPORTES =====
def _pdf_table(doc, els, headers, rows_data, col_widths, styles):
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    data = [headers] + rows_data
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d47a1")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
    ]))
    els.append(t)


def generar_pdf_reporte_ventas(db):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    data = get_reporte_ventas(db)
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=letter); styles = getSampleStyleSheet(); els = []
    els.append(Paragraph("MACUIN - Reporte de Ventas", styles["Title"]))
    els.append(Paragraph(f"Total Ventas: ${data['total_ventas']:,.2f}   |   Total Pedidos: {data['total_pedidos']}", styles["Normal"]))
    els.append(Spacer(1, 12))
    els.append(Paragraph("Clientes Top", styles["Heading2"]))
    _pdf_table(doc, els, ["Razón Social", "Pedidos", "Monto Total"],
               [[r["razon_social"] or "", str(r["total_pedidos"]), f"${r['monto_total']:,.2f}"] for r in data["clientes_top"]],
               [200, 80, 100], styles)
    els.append(Spacer(1, 12))
    els.append(Paragraph("Productos Más Vendidos", styles["Heading2"]))
    _pdf_table(doc, els, ["SKU", "Nombre", "Unidades Vendidas"],
               [[r["sku"], r["nombre"], str(r["total_vendido"])] for r in data["productos_top"]],
               [80, 240, 100], styles)
    doc.build(els); buf.seek(0); return buf


def generar_pdf_reporte_inventario(db):
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    data = get_reporte_inventario(db)
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=landscape(letter)); styles = getSampleStyleSheet(); els = []
    els.append(Paragraph("MACUIN - Reporte de Inventario", styles["Title"]))
    els.append(Paragraph(f"Total productos: {data['total_productos']}   |   Stock bajo: {data['productos_stock_bajo']}", styles["Normal"]))
    els.append(Spacer(1, 12))
    _pdf_table(doc, els, ["SKU", "Nombre", "Marca", "Categoría", "Stock Actual", "Stock Mínimo", "Ubicación", "Estatus"],
               [[r["sku"], r["nombre"], r["marca"] or "", r["categoria"], str(r["stock_actual"]),
                 str(r["stock_minimo"]), r["ubicacion"], r["estatus"]] for r in data["detalle"]],
               [55, 160, 70, 80, 55, 55, 60, 45], styles)
    doc.build(els); buf.seek(0); return buf


def generar_pdf_reporte_pedidos(db):
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    data = get_reporte_pedidos(db)
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=landscape(letter)); styles = getSampleStyleSheet(); els = []
    els.append(Paragraph("MACUIN - Reporte de Pedidos", styles["Title"]))
    els.append(Paragraph(f"Total pedidos: {data['total_pedidos']}   |   Monto total: ${data['monto_total']:,.2f}", styles["Normal"]))
    els.append(Spacer(1, 12))
    _pdf_table(doc, els, ["Folio", "Cliente", "Estatus", "Subtotal", "IVA", "Total", "Fecha"],
               [[r["folio"], r["cliente"], r["estatus"], f"${r['subtotal']:,.2f}",
                 f"${r['impuesto']:,.2f}", f"${r['total']:,.2f}", r["fecha"]] for r in data["detalle"]],
               [80, 160, 70, 70, 60, 70, 60], styles)
    doc.build(els); buf.seek(0); return buf


def generar_pdf_reporte_clientes(db):
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    data = get_reporte_clientes(db)
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=landscape(letter)); styles = getSampleStyleSheet(); els = []
    els.append(Paragraph("MACUIN - Reporte de Clientes", styles["Title"]))
    els.append(Paragraph(f"Total clientes: {data['total_clientes']}", styles["Normal"]))
    els.append(Spacer(1, 12))
    _pdf_table(doc, els, ["Nombre", "Email", "Razón Social", "RFC", "Ciudad", "Estado", "Pedidos", "Monto Total"],
               [[r["nombre"], r["email"], r["razon_social"], r["rfc"], r["ciudad"],
                 r["estado"], str(r["total_pedidos"]), f"${r['monto_total']:,.2f}"] for r in data["detalle"]],
               [90, 120, 110, 80, 70, 70, 50, 80], styles)
    doc.build(els); buf.seek(0); return buf


# ===== GENERADORES XLSX REPORTES =====
def _generar_xlsx(titulo, headers, rows):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = titulo
    ws.append([titulo]); ws["A1"].font = Font(bold=True, size=14)
    ws.append([])
    ws.append(headers)
    header_fill = PatternFill("solid", fgColor="0D47A1")
    for cell in ws[3]:
        cell.font = Font(bold=True, color="FFFFFF"); cell.fill = header_fill; cell.alignment = Alignment(horizontal="center")
    for row in rows:
        ws.append(row)
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = max(len(str(c.value or "")) for c in col) + 4
    buf = io.BytesIO(); wb.save(buf); buf.seek(0); return buf


def generar_xlsx_reporte_ventas(db):
    data = get_reporte_ventas(db)
    return _generar_xlsx("Reporte de Ventas",
        ["Razón Social", "Total Pedidos", "Monto Total"],
        [[r["razon_social"] or "", r["total_pedidos"], r["monto_total"]] for r in data["clientes_top"]])


def generar_xlsx_reporte_inventario(db):
    data = get_reporte_inventario(db)
    return _generar_xlsx("Reporte de Inventario",
        ["SKU", "Nombre", "Marca", "Categoría", "Stock Actual", "Stock Mínimo", "Ubicación", "Estatus"],
        [[r["sku"], r["nombre"], r["marca"], r["categoria"], r["stock_actual"], r["stock_minimo"], r["ubicacion"], r["estatus"]]
         for r in data["detalle"]])


def generar_xlsx_reporte_pedidos(db):
    data = get_reporte_pedidos(db)
    return _generar_xlsx("Reporte de Pedidos",
        ["Folio", "Cliente", "Estatus", "Subtotal", "IVA", "Total", "Fecha"],
        [[r["folio"], r["cliente"], r["estatus"], r["subtotal"], r["impuesto"], r["total"], r["fecha"]]
         for r in data["detalle"]])


def generar_xlsx_reporte_clientes(db):
    data = get_reporte_clientes(db)
    return _generar_xlsx("Reporte de Clientes",
        ["Nombre", "Email", "Razón Social", "RFC", "Ciudad", "Estado", "Total Pedidos", "Monto Total"],
        [[r["nombre"], r["email"], r["razon_social"], r["rfc"], r["ciudad"], r["estado"], r["total_pedidos"], r["monto_total"]]
         for r in data["detalle"]])


# ===== GENERADORES DOCX REPORTES =====
def _generar_docx(titulo, headers, rows):
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    doc.add_heading(titulo, 0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].paragraph_format.alignment = 1
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf


def generar_docx_reporte_ventas(db):
    data = get_reporte_ventas(db)
    return _generar_docx(f"Reporte de Ventas — Total: ${data['total_ventas']:,.2f}",
        ["Razón Social", "Total Pedidos", "Monto Total"],
        [[r["razon_social"] or "", r["total_pedidos"], f"${r['monto_total']:,.2f}"] for r in data["clientes_top"]])


def generar_docx_reporte_inventario(db):
    data = get_reporte_inventario(db)
    return _generar_docx(f"Reporte de Inventario — {data['total_productos']} productos",
        ["SKU", "Nombre", "Marca", "Categoría", "Stock Actual", "Stock Mínimo", "Ubicación", "Estatus"],
        [[r["sku"], r["nombre"], r["marca"], r["categoria"], r["stock_actual"], r["stock_minimo"], r["ubicacion"], r["estatus"]]
         for r in data["detalle"]])


def generar_docx_reporte_pedidos(db):
    data = get_reporte_pedidos(db)
    return _generar_docx(f"Reporte de Pedidos — Total: ${data['monto_total']:,.2f}",
        ["Folio", "Cliente", "Estatus", "Subtotal", "IVA", "Total", "Fecha"],
        [[r["folio"], r["cliente"], r["estatus"], f"${r['subtotal']:,.2f}", f"${r['impuesto']:,.2f}", f"${r['total']:,.2f}", r["fecha"]]
         for r in data["detalle"]])


def generar_docx_reporte_clientes(db):
    data = get_reporte_clientes(db)
    return _generar_docx(f"Reporte de Clientes — {data['total_clientes']} clientes",
        ["Nombre", "Email", "Razón Social", "RFC", "Ciudad", "Estado", "Total Pedidos", "Monto Total"],
        [[r["nombre"], r["email"], r["razon_social"], r["rfc"], r["ciudad"], r["estado"], r["total_pedidos"], f"${r['monto_total']:,.2f}"]
         for r in data["detalle"]])


# ===== PDF =====
def generar_pdf_pedido(db, pid):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    p = db.query(Pedido).filter(Pedido.id == pid).first()
    if not p: raise HTTPException(404, "Pedido no encontrado")
    buf = io.BytesIO(); doc = SimpleDocTemplate(buf, pagesize=letter); styles = getSampleStyleSheet(); els = []
    els.append(Paragraph("MACUIN - Autopartes Automotrices", styles["Title"]))
    els.append(Spacer(1, 12))
    els.append(Paragraph(f"Pedido: {p.folio}", styles["Heading2"]))
    els.append(Paragraph(f"Estatus: {p.estatus.upper()} | Fecha: {p.created_at.strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    if p.cliente and p.cliente.usuario:
        els.append(Paragraph(f"Cliente: {p.cliente.razon_social or ''} — {p.cliente.usuario.nombre} {p.cliente.usuario.apellido}", styles["Normal"]))
    els.append(Spacer(1, 20))
    data = [["SKU", "Producto", "Cant.", "P.Unit.", "Subtotal"]]
    for d in p.detalles:
        data.append([d.autoparte.sku, d.autoparte.nombre, str(d.cantidad), f"${d.precio_unitario:,.2f}", f"${d.subtotal:,.2f}"])
    t = Table(data, colWidths=[70, 200, 50, 80, 80])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d47a1")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                           ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 9),
                           ("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white])]))
    els.append(t); els.append(Spacer(1, 12))
    t2 = Table([["", "", "", "Subtotal:", f"${p.subtotal:,.2f}"], ["", "", "", "IVA 16%:", f"${p.impuesto:,.2f}"],
                ["", "", "", "TOTAL:", f"${p.total:,.2f}"]], colWidths=[70, 200, 50, 80, 80])
    t2.setStyle(TableStyle([("ALIGN", (3, 0), (-1, -1), "RIGHT"), ("FONTNAME", (3, 2), (-1, 2), "Helvetica-Bold")]))
    els.append(t2)
    if p.notas: els.append(Spacer(1, 20)); els.append(Paragraph(f"Notas: {p.notas}", styles["Normal"]))
    doc.build(els); buf.seek(0)
    return buf
